import argparse
import json
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

INK = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(89, 89, 89)
LIGHT_FILL = "F4F6F9"
RULE = "D7DBE2"


def export_tribble_docx_notes(
    data_dir: Path,
    output_dir: Path,
    run_ids: list[str] | None = None,
) -> dict[str, object]:
    runs = select_tribble_runs(data_dir, run_ids)
    output_dir.mkdir(parents=True, exist_ok=True)
    exported: list[dict[str, str]] = []

    for run in runs:
        title = note_title(run)
        path = output_dir / f"{safe_filename(title)}.docx"
        build_meeting_note(run).save(path)
        exported.append(
            {
                "run_id": str(run["run_id"]),
                "tribble_meeting_id": str(run["tribble_meeting_id"]),
                "title": title,
                "path": str(path.resolve()),
            }
        )

    return {
        "format": "docx",
        "output_dir": str(output_dir.resolve()),
        "exported": exported,
    }


def build_meeting_note(run: dict[str, Any]) -> Document:
    intelligence = run["intelligence"]
    document = Document()
    configure_document(document)
    add_masthead(document, run)

    add_prose_section(
        document,
        "Executive Summary",
        clean_text(intelligence.get("executive_summary")) or "None captured.",
        callout=True,
    )
    add_list_section(
        document,
        "Attendees",
        [
            attendee_display(attendee)
            for attendee in intelligence.get("attendees", [])
        ],
        always=True,
    )
    add_list_section(document, "Key Points", intelligence.get("key_points", []), always=True)
    add_list_section(document, "Decisions", intelligence.get("decisions", []))
    action_items = intelligence.get("action_items", [])
    if action_items:
        document.add_page_break()
    add_action_items(document, action_items)
    add_list_section(
        document,
        "Risks / Blockers",
        intelligence.get("risks_or_blockers", []),
    )
    add_list_section(document, "Open Questions", intelligence.get("open_questions", []))
    add_list_section(
        document,
        "Customer Pain Points",
        intelligence.get("customer_pain_points", []),
    )
    add_list_section(
        document,
        "Business Outcomes",
        intelligence.get("business_outcomes", []),
    )
    add_list_section(
        document,
        "Technical Requirements",
        intelligence.get("technical_requirements", []),
    )
    add_list_section(
        document,
        "Competitors Mentioned",
        intelligence.get("competitors_mentioned", []),
    )
    add_prose_section(
        document,
        "Recommended Next Step",
        clean_text(intelligence.get("proposed_next_step")) or "None captured.",
    )
    add_salesforce_match(document, run.get("salesforce_match", {}))
    add_email_draft(document, intelligence.get("follow_up_email_draft"))
    return document


def configure_document(document: Document) -> None:
    section = document.sections[0]
    document.settings.odd_and_even_pages_header_footer = False
    section.different_first_page_header_footer = False
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    set_style_font(normal, "Calibri", 11, RGBColor(0, 0, 0))
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    configure_heading(styles["Heading 1"], 16, BLUE, 18, 10)
    configure_heading(styles["Heading 2"], 13, BLUE, 14, 7)
    configure_heading(styles["Heading 3"], 12, DARK_BLUE, 10, 5)

    document.core_properties.title = "Tribble Meeting Notes"
    document.core_properties.author = "Meeting Notes Agent"
    document.core_properties.subject = "Meeting recap generated from Tribble Scribe"

    header = section.header.paragraphs[0]
    header.text = "MEETING NOTES"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    set_paragraph_runs(header, "Calibri", 8.5, GRAY, bold=True)

    footer = section.footer.paragraphs[0]
    footer.text = "Generated from Tribble Scribe"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    set_paragraph_runs(footer, "Calibri", 8.5, GRAY)


def add_masthead(document: Document, run: dict[str, Any]) -> None:
    intelligence = run["intelligence"]
    kicker = document.add_paragraph()
    kicker.paragraph_format.space_after = Pt(3)
    kicker_run = kicker.add_run("TRIBBLE MEETING NOTES")
    set_run_font(kicker_run, "Calibri", 9, BLUE, bold=True)

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)
    title.paragraph_format.keep_with_next = True
    title_run = title.add_run(clean_text(intelligence.get("meeting_title")) or "Meeting")
    set_run_font(title_run, "Calibri", 23, INK, bold=True)

    metadata = [
        ("Date", clean_text(intelligence.get("meeting_date")) or "Not captured"),
        ("Source", "Tribble Scribe"),
        ("Run ID", str(run["run_id"])),
    ]
    for label, value in metadata:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.1
        label_run = paragraph.add_run(f"{label}: ")
        set_run_font(label_run, "Calibri", 10, GRAY, bold=True)
        value_run = paragraph.add_run(value)
        set_run_font(value_run, "Calibri", 10, GRAY)

    rule = document.add_paragraph()
    rule.paragraph_format.space_before = Pt(8)
    rule.paragraph_format.space_after = Pt(2)
    add_bottom_border(rule, RULE, 8)


def add_prose_section(
    document: Document,
    heading: str,
    value: str,
    callout: bool = False,
) -> None:
    add_heading(document, heading)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.keep_together = True
    if callout:
        paragraph.paragraph_format.left_indent = Inches(0.14)
        paragraph.paragraph_format.right_indent = Inches(0.14)
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(8)
        add_paragraph_shading(paragraph, LIGHT_FILL)
    run = paragraph.add_run(clean_text(value))
    set_run_font(run, "Calibri", 11, RGBColor(0, 0, 0))


def add_list_section(
    document: Document,
    heading: str,
    values: list[str],
    always: bool = False,
) -> None:
    cleaned = [clean_text(value) for value in values if clean_text(value)]
    if not cleaned and not always:
        return
    add_heading(document, heading)
    if not cleaned:
        document.add_paragraph("None captured.")
        return
    for value in cleaned:
        add_bullet(document, value)


def add_action_items(document: Document, items: list[dict[str, Any]]) -> None:
    add_heading(document, "Action Items")
    if not items:
        document.add_paragraph("None captured.")
        return
    for item in items:
        paragraph = add_bullet(document)
        owner = clean_text(item.get("owner")) or "Unassigned"
        owner_run = paragraph.add_run(f"{owner}: ")
        set_run_font(owner_run, "Calibri", 11, RGBColor(0, 0, 0), bold=True)
        task_run = paragraph.add_run(clean_text(item.get("task")))
        set_run_font(task_run, "Calibri", 11, RGBColor(0, 0, 0))
        due_date = clean_text(item.get("due_date"))
        if due_date:
            due_run = paragraph.add_run(f"  Due: {due_date}")
            set_run_font(due_run, "Calibri", 10, GRAY, italic=True)


def add_salesforce_match(document: Document, match: dict[str, Any]) -> None:
    account = clean_text(match.get("account_name"))
    opportunity = clean_text(match.get("opportunity_name"))
    reasons = [clean_text(value) for value in match.get("reasons", []) if value]
    status = clean_text(match.get("match_status") or "no_match").replace("_", " ").title()
    section_paragraphs = [add_heading(document, "Salesforce Match")]
    section_paragraphs.append(add_bullet(document, f"Status: {status}"))
    section_paragraphs.append(add_bullet(document, f"Account: {account or 'No match'}"))
    section_paragraphs.append(
        add_bullet(document, f"Opportunity: {opportunity or 'No match'}")
    )
    section_paragraphs.append(
        add_bullet(document, f"Confidence: {match.get('confidence', 0)}")
    )
    for reason in reasons:
        section_paragraphs.append(add_bullet(document, f"Reason: {reason}"))
    for alternative in match.get("alternatives", [])[:3]:
        candidate_account = clean_text(alternative.get("account_name")) or "Unknown Account"
        candidate_opportunity = (
            clean_text(alternative.get("opportunity_name")) or "Unknown Opportunity"
        )
        section_paragraphs.append(
            add_bullet(
                document,
                "Candidate: "
                f"{candidate_account} / {candidate_opportunity} "
                f"({alternative.get('confidence', 0)})",
            )
        )
    for paragraph in section_paragraphs[:-1]:
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.keep_together = True


def add_email_draft(document: Document, value: str | None) -> None:
    add_heading(document, "Follow-Up Email Draft")
    lines = clean_text(value).splitlines() if value else []
    if not lines:
        document.add_paragraph("None captured.")
        return
    for line in lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.2)
        paragraph.paragraph_format.right_indent = Inches(0.2)
        paragraph.paragraph_format.space_after = Pt(4 if line else 2)
        paragraph.paragraph_format.line_spacing = 1.15
        run = paragraph.add_run(line)
        set_run_font(run, "Calibri", 10.5, RGBColor(35, 35, 35))


def add_heading(document: Document, text: str):
    heading = document.add_heading(text, level=1)
    heading.paragraph_format.keep_with_next = True
    return heading


def add_bullet(document: Document, value: str | None = None):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    # A small leading gap keeps adjacent numbered paragraphs visually distinct in
    # Word's PDF renderer, including when a list crosses a page boundary.
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.widow_control = True
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = False
    if value:
        run = paragraph.add_run(clean_text(value))
        set_run_font(run, "Calibri", 11, RGBColor(0, 0, 0))
    return paragraph


def configure_bullet_numbering(document: Document) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "bullet")
    level.append(number_format)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "\u2022")
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)

    paragraph_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    paragraph_properties.append(tabs)
    indentation = OxmlElement("w:ind")
    indentation.set(qn("w:left"), "540")
    indentation.set(qn("w:hanging"), "270")
    paragraph_properties.append(indentation)
    level.append(paragraph_properties)

    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:hint"), "default")
    run_properties.append(fonts)
    level.append(run_properties)
    abstract.append(level)
    numbering.append(abstract)

    return abstract_id


def create_numbering_instance(document: Document, abstract_id: int) -> int:
    """Create a fresh list instance so Word renders every bullet reliably."""
    numbering = document.part.numbering_part.element
    num_ids = [
        int(element.get(qn("w:numId")))
        for element in numbering.findall(qn("w:num"))
    ]
    num_id = max(num_ids, default=0) + 1

    concrete = OxmlElement("w:num")
    concrete.set(qn("w:numId"), str(num_id))
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), str(abstract_id))
    concrete.append(abstract_reference)
    numbering.append(concrete)
    return num_id


def configure_heading(style, size: float, color: RGBColor, before: float, after: float) -> None:
    set_style_font(style, "Calibri", size, color)
    style.font.bold = True
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.keep_with_next = True


def set_style_font(style, name: str, size: float, color: RGBColor) -> None:
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    style.font.size = Pt(size)
    style.font.color.rgb = color


def set_run_font(
    run,
    name: str,
    size: float,
    color: RGBColor,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_runs(
    paragraph,
    name: str,
    size: float,
    color: RGBColor,
    bold: bool | None = None,
) -> None:
    for run in paragraph.runs:
        set_run_font(run, name, size, color, bold=bold)


def add_bottom_border(paragraph, color: str, size: int) -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        properties.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_paragraph_shading(paragraph, fill: str) -> None:
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def attendee_display(attendee: dict[str, Any]) -> str:
    name = clean_text(attendee.get("name")) or "Unknown"
    email = clean_text(attendee.get("email"))
    return f"{name} ({email})" if email else name


def note_title(run: dict[str, Any]) -> str:
    intelligence = run["intelligence"]
    meeting_date = clean_text(intelligence.get("meeting_date"))
    meeting_title = clean_text(intelligence.get("meeting_title")) or "Meeting"
    return f"{meeting_date} - {meeting_title}" if meeting_date else meeting_title


def safe_filename(value: str) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*]', "-", value)
    return re.sub(r"\s+", " ", cleaned).strip(" .")


def clean_text(value: Any) -> str:
    if not value:
        return ""
    text = str(value)
    if not any(marker in text for marker in ("Ã", "Â", "â€", "â€™", "â€”", "â€“")):
        return text
    try:
        return text.encode("cp1252").decode("utf-8")
    except (UnicodeEncodeError, UnicodeDecodeError):
        return text


def select_tribble_runs(
    data_dir: Path,
    run_ids: list[str] | None,
) -> list[dict[str, Any]]:
    paths = sorted(data_dir.glob("*.json"))
    by_id: dict[str, dict[str, Any]] = {}
    for path in paths:
        run = json.loads(path.read_text(encoding="utf-8"))
        if run.get("tribble_meeting_id"):
            by_id[str(run["run_id"])] = run

    if run_ids:
        missing = [run_id for run_id in run_ids if run_id not in by_id]
        if missing:
            raise ValueError("Tribble run not found: " + ", ".join(missing))
        return [by_id[run_id] for run_id in run_ids]

    return sorted(by_id.values(), key=lambda run: str(run["created_at"]))


def main() -> None:
    parser = argparse.ArgumentParser(
        prog="sharepoint-docx-export",
        description="Export Tribble meeting-note runs as Word documents.",
    )
    parser.add_argument(
        "--data-dir",
        type=Path,
        default=Path("./data/runs"),
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=Path("./out/sharepoint-docx"),
    )
    parser.add_argument("--run-id", action="append", dest="run_ids")
    args = parser.parse_args()
    result = export_tribble_docx_notes(
        data_dir=args.data_dir,
        output_dir=args.output_dir,
        run_ids=args.run_ids,
    )
    print(json.dumps(result, indent=2))


if __name__ == "__main__":
    main()
